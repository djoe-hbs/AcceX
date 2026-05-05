import os
import shutil
import tempfile
import zipfile
from pathlib import Path

from django.conf import settings
from django.core.files import File
from django.db import transaction

from pypdf import PdfReader
from docx import Document
from docx.oxml.ns import qn
from openpyxl import load_workbook

from core.work.models import WorkBatch, WorkFile, normalize_work_relative_path


SUPPORTED_EXCEL_EXTENSIONS = {".xlsx", ".xlsm", ".xltx", ".xltm"}
SUPPORTED_DOCX_EXTENSIONS = {".docx"}


def _safe_extract_zip(zip_path: Path, destination: Path):
    destination.mkdir(parents=True, exist_ok=True)

    with zipfile.ZipFile(zip_path, "r") as zip_ref:
        for member in zip_ref.infolist():
            normalized_name = normalize_work_relative_path(member.filename)
            if not normalized_name:
                continue

            member_path = destination / Path(*normalized_name.split("/"))
            resolved_member_path = member_path.resolve()
            resolved_destination = destination.resolve()

            try:
                resolved_member_path.relative_to(resolved_destination)
            except ValueError:
                continue

            if member.is_dir():
                resolved_member_path.mkdir(parents=True, exist_ok=True)
                continue

            resolved_member_path.parent.mkdir(parents=True, exist_ok=True)
            with zip_ref.open(member) as src, open(resolved_member_path, "wb") as dst:
                shutil.copyfileobj(src, dst)


def _unique_extract_dir_for_zip(zip_file: Path):
    base_dir = zip_file.parent / zip_file.stem
    candidate = base_dir
    suffix = 1

    while candidate.exists():
        candidate = zip_file.parent / f"{zip_file.stem}_{suffix}"
        suffix += 1

    return candidate


def _extract_nested_archives(root_dir: Path):
    processed = set()

    while True:
        pending_archives = [
            path for path in root_dir.rglob("*.zip")
            if path.is_file() and path not in processed
        ]

        if not pending_archives:
            break

        for archive in pending_archives:
            extract_dir = _unique_extract_dir_for_zip(archive)
            _safe_extract_zip(archive, extract_dir)
            processed.add(archive)


def _classify_file(path: Path):
    ext = path.suffix.lower()

    if ext == ".pdf":
        return WorkFile.FileType.PDF, WorkFile.CountType.PAGE
    if ext in SUPPORTED_DOCX_EXTENSIONS:
        return WorkFile.FileType.DOCX, WorkFile.CountType.PAGE
    if ext in SUPPORTED_EXCEL_EXTENSIONS:
        return WorkFile.FileType.EXCEL, WorkFile.CountType.ROW
    if ext == ".zip":
        return WorkFile.FileType.ZIP, WorkFile.CountType.NONE

    return WorkFile.FileType.OTHER, WorkFile.CountType.LINE


def _count_pdf_pages(path: Path):
    reader = PdfReader(str(path))
    return len(reader.pages)


def _count_docx_pages(path: Path):
    document = Document(str(path))

    if not document.paragraphs:
        return 0

    page_breaks = 0
    for paragraph in document.paragraphs:
        for run in paragraph.runs:
            run_element = run._element

            for break_element in run_element.findall(".//w:br", run_element.nsmap):
                if break_element.get(qn("w:type")) == "page":
                    page_breaks += 1

            page_breaks += len(run_element.findall(".//w:lastRenderedPageBreak", run_element.nsmap))

    return max(1, page_breaks + 1)


def _count_excel_rows(path: Path):
    workbook = load_workbook(filename=str(path), read_only=True, data_only=True)
    row_count = 0

    try:
        for sheet in workbook.worksheets:
            for row in sheet.iter_rows(values_only=True):
                if any(cell not in (None, "") for cell in row):
                    row_count += 1
    finally:
        workbook.close()

    return row_count


def _count_file_lines(path: Path):
    with open(path, "rb") as handle:
        return sum(1 for _ in handle)


def _calculate_count(path: Path, file_type):
    try:
        if file_type == WorkFile.FileType.PDF:
            return _count_pdf_pages(path)
        if file_type == WorkFile.FileType.DOCX:
            return _count_docx_pages(path)
        if file_type == WorkFile.FileType.EXCEL:
            return _count_excel_rows(path)
        if file_type == WorkFile.FileType.OTHER:
            return _count_file_lines(path)
    except Exception:
        return None

    return None


def _build_file_tree(batch: WorkBatch, extraction_root: Path):
    WorkFile.objects.filter(batch=batch).delete()

    node_lookup = {}
    total_files = 0
    total_directories = 0

    for current_root, dirs, files in sorted(os.walk(extraction_root)):
        current_root = Path(current_root)
        relative_root = current_root.relative_to(extraction_root)

        dirs.sort()
        files.sort()

        for directory in dirs:
            relative_path = normalize_work_relative_path(
                (relative_root / directory).as_posix() if str(relative_root) != "." else directory
            )
            parent_key = Path(relative_path).parent.as_posix()
            parent = node_lookup.get(parent_key) if parent_key != "." else None

            node = WorkFile.objects.create(
                batch=batch,
                parent=parent,
                name=directory,
                relative_path=relative_path,
                depth=len(Path(relative_path).parts),
                is_directory=True,
                file_extension=None,
                file_type=WorkFile.FileType.DIRECTORY,
                count_type=WorkFile.CountType.NONE,
                count=None,
                size_bytes=0,
            )
            node_lookup[relative_path] = node
            total_directories += 1

        for file_name in files:
            absolute_file = current_root / file_name
            relative_path = normalize_work_relative_path(
                (relative_root / file_name).as_posix() if str(relative_root) != "." else file_name
            )
            parent_key = Path(relative_path).parent.as_posix()
            parent = node_lookup.get(parent_key) if parent_key != "." else None

            file_type, count_type = _classify_file(absolute_file)
            count = _calculate_count(absolute_file, file_type)

            node = WorkFile.objects.create(
                batch=batch,
                parent=parent,
                name=file_name,
                relative_path=relative_path,
                depth=len(Path(relative_path).parts),
                is_directory=False,
                file_extension=absolute_file.suffix.lower() or None,
                file_type=file_type,
                count_type=count_type,
                count=count,
                size_bytes=absolute_file.stat().st_size,
            )
            with absolute_file.open("rb") as stream:
                node.source_file.save(Path(relative_path).name, File(stream), save=False)
            node.save(update_fields=["source_file", "updated"])
            node_lookup[relative_path] = node
            total_files += 1

    batch.total_files = total_files
    batch.total_directories = total_directories


def process_work_batch(batch: WorkBatch):
    extraction_root = Path(settings.WORK_EXTRACTION_ROOT) / batch.public_id.hex
    extraction_root.mkdir(parents=True, exist_ok=True)

    batch.status = WorkBatch.Status.PROCESSING
    batch.error_message = None
    batch.extraction_root = batch.public_id.hex
    batch.save(update_fields=["status", "error_message", "extraction_root", "updated"])

    try:
        # Download source archive to a local temp file (supports both local and S3 storage)
        with tempfile.NamedTemporaryFile(suffix=".zip", delete=False) as tmp:
            tmp_path = Path(tmp.name)
            with batch.source_archive.open("rb") as src:
                shutil.copyfileobj(src, tmp)

        try:
            _safe_extract_zip(tmp_path, extraction_root)
        finally:
            tmp_path.unlink(missing_ok=True)

        _extract_nested_archives(extraction_root)

        with transaction.atomic():
            _build_file_tree(batch, extraction_root)
            batch.status = WorkBatch.Status.READY
            batch.error_message = None
            batch.save(update_fields=["status", "error_message", "total_files", "total_directories", "updated"])
    except Exception as exc:
        batch.status = WorkBatch.Status.FAILED
        batch.error_message = str(exc)
        batch.save(update_fields=["status", "error_message", "updated"])
        raise
    finally:
        shutil.rmtree(extraction_root, ignore_errors=True)
