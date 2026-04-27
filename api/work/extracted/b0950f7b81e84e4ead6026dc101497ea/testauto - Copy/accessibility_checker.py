# """
# Word Document Accessibility Checker & Auto-Fixer
# Based on Word Document Accessibility Guidelines
# """

# import zipfile
# import shutil
# import re
# import os
# import copy
# from pathlib import Path
# from lxml import etree
# from docx import Document
# from docx.oxml.ns import qn
# from docx.enum.text import WD_ALIGN_PARAGRAPH
# from docx.oxml import OxmlElement


# # ─────────────────────── XML Namespaces ──────────────────────
# NS = {
#     "w":   "http://schemas.openxmlformats.org/wordprocessingml/2006/main",
#     "r":   "http://schemas.openxmlformats.org/officeDocument/2006/relationships",
#     "wp":  "http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing",
#     "a":   "http://schemas.openxmlformats.org/drawingml/2006/main",
#     "pic": "http://schemas.openxmlformats.org/drawingml/2006/picture",
#     "mc":  "http://schemas.openxmlformats.org/markup-compatibility/2006",
# }


# # ─────────────────────── Issue Registry ──────────────────────
# class Issue:
#     def __init__(self, rule_id: str, severity: str, location: str, description: str, fix_applied: bool = False):
#         self.rule_id = rule_id
#         self.severity = severity      # "ERROR" | "WARNING" | "INFO"
#         self.location = location
#         self.description = description
#         self.fix_applied = fix_applied

#     def __repr__(self):
#         marker = "✔ FIXED" if self.fix_applied else "✗ MANUAL"
#         return f"[{self.severity}][{self.rule_id}] {self.location}: {self.description} — {marker}"


# # ═════════════════════════════════════════════════════════════
# #  AccessibilityChecker
# # ═════════════════════════════════════════════════════════════
# class AccessibilityChecker:
#     """
#     Checks and (where possible) auto-fixes a .docx file against the
#     Word Document Accessibility Guidelines.
#     """

#     # Heading style names recognised by Word
#     HEADING_STYLES = {
#         "heading 1": 1, "heading 2": 2, "heading 3": 3,
#         "heading 4": 4, "heading 5": 5, "heading 6": 6,
#     }

#     # Plain-text math symbols allowed outside an equation object
#     ALLOWED_PLAIN_MATH_CHARS = set("+=<>$%÷±≠≈≤≥∞−×Δ")

#     def __init__(self, input_path: str, output_path: str | None = None, auto_fix: bool = True):
#         self.input_path = Path(input_path)
#         self.output_path = Path(output_path) if output_path else self.input_path.with_stem(
#             self.input_path.stem + "_accessible"
#         )
#         self.auto_fix = auto_fix
#         self.issues: list[Issue] = []
#         self.doc: Document | None = None

#     # ─────────────────── Public entry point ──────────────────
#     def run(self) -> list[Issue]:
#         """Run all checks and fixes. Returns list of Issue objects."""
#         print(f"\n{'='*60}")
#         print(f"  Word Accessibility Checker")
#         print(f"  Input : {self.input_path}")
#         print(f"  Output: {self.output_path}")
#         print(f"  Mode  : {'Auto-fix ON' if self.auto_fix else 'Audit only'}")
#         print(f"{'='*60}\n")

#         self.doc = Document(str(self.input_path))

#         # ── Run all rule groups ──
#         self._check_headings()
#         self._check_empty_paragraphs()
#         self._check_double_spaces()
#         self._check_underlines()
#         self._check_lists()
#         self._check_tables()
#         self._check_links()
#         self._check_images()
#         self._check_equations()
#         self._check_metadata()
#         self._check_color_contrast()
#         self._check_page_breaks()
#         self._check_bookmarks()

#         # ── Save fixed document ──
#         if self.auto_fix:
#             self.doc.save(str(self.output_path))
#             print(f"\n✅ Fixed document saved to: {self.output_path}")

#         # ── Print summary ──
#         self._print_summary()
#         return self.issues

#     # ═══════════════════════════════════════════════════════
#     #  RULE CHECKS
#     # ═══════════════════════════════════════════════════════

#     # ── R1: Heading hierarchy ────────────────────────────────
#     def _check_headings(self):
#         """
#         • Only one H1 per document.
#         • No skipped levels (H1 → H3 without H2 is invalid).
#         • Font sizes must not be manually overridden on headings.
#         """
#         heading_levels = []
#         h1_count = 0

#         for para in self.doc.paragraphs:
#             style_name = para.style.name.lower()
#             if style_name not in self.HEADING_STYLES:
#                 continue
#             level = self.HEADING_STYLES[style_name]
#             heading_levels.append((level, para))

#             if level == 1:
#                 h1_count += 1

#         # Check single H1
#         if h1_count == 0:
#             self.issues.append(Issue("H1-MISSING", "WARNING", "Document",
#                                      "No H1 (Heading 1) found. Every document should have exactly one H1."))
#         elif h1_count > 1:
#             self.issues.append(Issue("H1-MULTIPLE", "ERROR", "Document",
#                                      f"Found {h1_count} Heading-1 paragraphs. Only one H1 is allowed."))

#         # Check level skipping
#         prev_level = 0
#         for level, para in heading_levels:
#             if level > prev_level + 1 and prev_level != 0:
#                 loc = f"Para: '{para.text[:60]}...'" if len(para.text) > 60 else f"Para: '{para.text}'"
#                 self.issues.append(Issue("HEADING-SKIP", "ERROR", loc,
#                                          f"Heading level skipped: H{prev_level} → H{level}"))
#             prev_level = level

#         # Check font-size overrides on headings
#         for level, para in heading_levels:
#             for run in para.runs:
#                 sz_el = run._r.find(qn("w:rPr") + "/" + qn("w:sz"), NS)
#                 # Walk direct rPr child
#                 rpr = run._r.find(qn("w:rPr"))
#                 if rpr is not None:
#                     sz = rpr.find(qn("w:sz"))
#                     if sz is not None:
#                         loc = f"H{level}: '{para.text[:50]}'"
#                         self.issues.append(Issue("HEADING-FONTSIZE", "WARNING", loc,
#                                                   "Manual font size found on heading. Remove to keep original heading size."))
#                         if self.auto_fix:
#                             rpr.remove(sz)
#                         break

#         print(f"  [R1] Headings checked. H1 count={h1_count}, levels checked.")

#     # ── R2: Empty paragraphs ────────────────────────────────
#     def _check_empty_paragraphs(self):
#         """
#         No empty paragraphs for spacing — use paragraph spacing instead.
#         Auto-fix: remove consecutive empty paragraphs; flag single empty ones.
#         """
#         paragraphs = self.doc.paragraphs
#         i = 0
#         removed = 0
#         while i < len(paragraphs) - 1:
#             curr = paragraphs[i]
#             nxt  = paragraphs[i + 1]
#             if curr.text.strip() == "" and nxt.text.strip() == "":
#                 loc = f"Para index ~{i}"
#                 self.issues.append(Issue("EMPTY-PARA", "ERROR", loc,
#                                           "Consecutive empty paragraphs used for spacing.",
#                                           fix_applied=self.auto_fix))
#                 if self.auto_fix:
#                     p_el = curr._element
#                     p_el.getparent().remove(p_el)
#                     removed += 1
#                     # Don't advance; re-check same index with new list
#                     paragraphs = self.doc.paragraphs
#                     continue
#             i += 1

#         print(f"  [R2] Empty paragraphs: {removed} removed.")

#     # ── R3: Double spaces ───────────────────────────────────
#     def _check_double_spaces(self):
#         """Replace double (or more) spaces with single space in all runs."""
#         fixed = 0
#         for para in self.doc.paragraphs:
#             for run in para.runs:
#                 if "  " in run.text:
#                     self.issues.append(Issue("DOUBLE-SPACE", "WARNING",
#                                               f"Para: '{para.text[:50]}'",
#                                               "Multiple consecutive spaces found.",
#                                               fix_applied=self.auto_fix))
#                     if self.auto_fix:
#                         run.text = re.sub(r"  +", " ", run.text)
#                         fixed += 1

#         print(f"  [R3] Double spaces: {fixed} runs fixed.")

#     # ── R4: Underlines ──────────────────────────────────────
#     def _check_underlines(self):
#         """
#         No underlines except on hyperlinks. (Links MUST be underlined.)
#         """
#         fixed = 0
#         for para in self.doc.paragraphs:
#             # Collect hyperlink rIds so we can skip those runs
#             hyperlink_run_ids = set()
#             for hl in para._p.findall(".//" + qn("w:hyperlink")):
#                 for r in hl.findall(".//" + qn("w:r")):
#                     hyperlink_run_ids.add(id(r))

#             for run in para.runs:
#                 if id(run._r) in hyperlink_run_ids:
#                     continue   # skip hyperlink runs
#                 if run.underline:
#                     self.issues.append(Issue("UNDERLINE", "ERROR",
#                                               f"Para: '{para.text[:50]}'",
#                                               f"Underline found on non-link text: '{run.text[:40]}'",
#                                               fix_applied=self.auto_fix))
#                     if self.auto_fix:
#                         run.underline = False
#                         fixed += 1

#         print(f"  [R4] Underlines: {fixed} non-link underlines removed.")

#     # ── R5: Lists ───────────────────────────────────────────
#     def _check_lists(self):
#         """
#         Detect paragraphs that look like lists but use manual numbering/bullets.
#         Patterns: "1. text", "a) text", "• text" (unicode bullet not via numPr).
#         """
#         MANUAL_LIST_RE = re.compile(
#             r"^(\s*(\d+[\.\)]\s|[a-zA-Z][\.\)]\s|[•·–\-]\s))"
#         )
#         warned = 0
#         for para in self.doc.paragraphs:
#             # Check if paragraph already uses Word numbering
#             num_pr = para._p.find(".//" + qn("w:numPr"))
#             if num_pr is not None:
#                 continue  # Proper list — OK

#             txt = para.text.strip()
#             if MANUAL_LIST_RE.match(txt):
#                 self.issues.append(Issue("MANUAL-LIST", "ERROR",
#                                           f"Para: '{txt[:60]}'",
#                                           "Manual list formatting detected. Use Home → Paragraph → Bullets/Numbering."))
#                 warned += 1

#         print(f"  [R5] Lists: {warned} manual-list paragraphs flagged (require manual fix).")

#     # ── R6: Tables ──────────────────────────────────────────
#     def _check_tables(self):
#         """
#         • No blank rows/columns.
#         • Header row should be marked (tblHeader).
#         • No layout-only tables (heuristic: single-cell).
#         """
#         for t_idx, table in enumerate(self.doc.tables):
#             loc = f"Table {t_idx + 1}"

#             # Check for blank rows
#             for r_idx, row in enumerate(table.rows):
#                 if all(cell.text.strip() == "" for cell in row.cells):
#                     self.issues.append(Issue("TABLE-BLANK-ROW", "ERROR",
#                                               f"{loc}, Row {r_idx + 1}",
#                                               "Completely blank row found in table.",
#                                               fix_applied=self.auto_fix))
#                     if self.auto_fix:
#                         tr = row._tr
#                         tr.getparent().remove(tr)

#             # Check for blank columns (re-read in case rows were removed)
#             if table.columns:
#                 num_cols = len(table.columns)
#                 for c_idx in range(num_cols):
#                     try:
#                         col_cells = [table.cell(r, c_idx) for r in range(len(table.rows))]
#                         if all(c.text.strip() == "" for c in col_cells):
#                             self.issues.append(Issue("TABLE-BLANK-COL", "ERROR",
#                                                       f"{loc}, Col {c_idx + 1}",
#                                                       "Completely blank column found in table."))
#                     except IndexError:
#                         pass

#             # Check header row is defined (tblHeader on first row's trPr)
#             if table.rows:
#                 first_tr = table.rows[0]._tr
#                 trPr = first_tr.find(qn("w:trPr"))
#                 tblHeader = trPr.find(qn("w:tblHeader")) if trPr is not None else None
#                 if tblHeader is None:
#                     self.issues.append(Issue("TABLE-NO-HEADER", "WARNING",
#                                               loc,
#                                               "Header row not marked as repeating header.",
#                                               fix_applied=self.auto_fix))
#                     if self.auto_fix:
#                         if trPr is None:
#                             trPr = OxmlElement("w:trPr")
#                             first_tr.insert(0, trPr)
#                         hdr = OxmlElement("w:tblHeader")
#                         trPr.append(hdr)

#             # Single-cell table = possible layout table
#             if len(table.rows) == 1 and len(table.columns) == 1:
#                 self.issues.append(Issue("TABLE-LAYOUT", "WARNING",
#                                           loc,
#                                           "Single-cell table may be used for layout rather than data."))

#         print(f"  [R6] Tables: {len(self.doc.tables)} tables checked.")

#     # ── R7: Links ───────────────────────────────────────────
#     def _check_links(self):
#         """
#         All hyperlinks must be underlined and must resolve (non-empty URL).
#         Internal links (cross-references) are checked for non-empty anchor.
#         """
#         fixed = 0
#         for para in self.doc.paragraphs:
#             for hl in para._p.findall(".//" + qn("w:hyperlink")):
#                 rId = hl.get(qn("r:id"))
#                 anchor = hl.get(qn("w:anchor"))

#                 # Check underline on all runs inside hyperlink
#                 for run_el in hl.findall(".//" + qn("w:r")):
#                     rpr = run_el.find(qn("w:rPr"))
#                     if rpr is None:
#                         rpr = OxmlElement("w:rPr")
#                         run_el.insert(0, rpr)
#                     u_el = rpr.find(qn("w:u"))
#                     if u_el is None:
#                         self.issues.append(Issue("LINK-NOT-UNDERLINED", "ERROR",
#                                                   f"Para: '{para.text[:50]}'",
#                                                   "Hyperlink run is not underlined.",
#                                                   fix_applied=self.auto_fix))
#                         if self.auto_fix:
#                             u = OxmlElement("w:u")
#                             u.set(qn("w:val"), "single")
#                             rpr.append(u)
#                             fixed += 1

#                 # Empty anchor = broken internal link
#                 if anchor is not None and anchor.strip() == "":
#                     self.issues.append(Issue("LINK-EMPTY-ANCHOR", "WARNING",
#                                               f"Para: '{para.text[:50]}'",
#                                               "Internal hyperlink has empty anchor."))

#         print(f"  [R7] Links: {fixed} link underlines added.")

#     # ── R8: Images ──────────────────────────────────────────
#     def _check_images(self):
#         """
#         • Every image must have non-empty alt text.
#         • Images must NOT be marked decorative.
#         • AI-generated alt text warning (heuristic: very long or generic).
#         """
#         img_count = 0
#         for para in self.doc.paragraphs:
#             for drawing in para._p.findall(".//" + qn("wp:inline")) + \
#                            para._p.findall(".//" + qn("wp:anchor")):
#                 img_count += 1
#                 loc = f"Image in para: '{para.text[:40]}'"

#                 # Alt text lives in wp:docPr/@descr
#                 docPr = drawing.find(qn("wp:docPr"))
#                 if docPr is None:
#                     self.issues.append(Issue("IMAGE-NO-ALTTEXT", "ERROR", loc,
#                                               "Image is missing alt text (wp:docPr not found)."))
#                     continue

#                 descr = docPr.get("descr", "").strip()
#                 hidden = docPr.get("hidden", "0")
#                 name   = docPr.get("name", "")

#                 # Decorative check
#                 if hidden == "1":
#                     self.issues.append(Issue("IMAGE-DECORATIVE", "ERROR", loc,
#                                               "Image is marked as decorative. Per guidelines, images must NOT be decorative.",
#                                               fix_applied=self.auto_fix))
#                     if self.auto_fix:
#                         docPr.set("hidden", "0")

#                 # Empty alt text
#                 if not descr:
#                     self.issues.append(Issue("IMAGE-NO-ALTTEXT", "ERROR", loc,
#                                               "Image has no alt text description."))

#                 # Heuristic: auto-generated AI alt text (Word's AI tends to produce very long descriptions)
#                 elif len(descr) > 300:
#                     self.issues.append(Issue("IMAGE-AI-ALTTEXT", "WARNING", loc,
#                                               "Alt text is very long — verify it wasn't auto-generated by Word's AI."))

#         print(f"  [R8] Images: {img_count} images checked.")

#     # ── R9: Equations ───────────────────────────────────────
#     def _check_equations(self):
#         """
#         • Equations must not be images.
#         • Equations must not be bold or italic.
#         • Plain-text math outside equation objects: only allowed symbols.
#         """
#         warned = 0
#         for para in self.doc.paragraphs:
#             # Check for Office Math runs (oMath)
#             omath = para._p.findall(".//{http://schemas.openxmlformats.org/officeDocument/2006/math}oMath")
#             for eq in omath:
#                 # Check bold/italic on math runs
#                 for r in eq.findall(".//" + qn("w:r")):
#                     rpr = r.find(qn("w:rPr"))
#                     if rpr is not None:
#                         b = rpr.find(qn("w:b"))
#                         i = rpr.find(qn("w:i"))
#                         if b is not None or i is not None:
#                             self.issues.append(Issue("EQ-BOLD-ITALIC", "ERROR",
#                                                       f"Para: '{para.text[:50]}'",
#                                                       "Equation contains bold or italic formatting — must be removed.",
#                                                       fix_applied=self.auto_fix))
#                             if self.auto_fix:
#                                 if b is not None:
#                                     rpr.remove(b)
#                                 if i is not None:
#                                     rpr.remove(i)
#                                 warned += 1

#         print(f"  [R9] Equations: {warned} bold/italic fixes applied.")

#     # ── R10: Metadata ────────────────────────────────────────
#     def _check_metadata(self):
#         """
#         Document must have a Title set. Author should be set.
#         """
#         props = self.doc.core_properties
#         if not props.title or props.title.strip() == "":
#             self.issues.append(Issue("META-NO-TITLE", "ERROR", "Document Properties",
#                                       "Document has no title set. Add via File → Info → Properties."))
#         if not props.author or props.author.strip() == "":
#             self.issues.append(Issue("META-NO-AUTHOR", "WARNING", "Document Properties",
#                                       "Document has no author set."))

#         print(f"  [R10] Metadata: title='{props.title}', author='{props.author}'.")

#     # ── R11: Color contrast ──────────────────────────────────
#     def _check_color_contrast(self):
#         """
#         Heuristic: warn if non-default font colours are found (manual review needed).
#         """
#         warned = 0
#         for para in self.doc.paragraphs:
#             for run in para.runs:
#                 rpr = run._r.find(qn("w:rPr"))
#                 if rpr is None:
#                     continue
#                 color_el = rpr.find(qn("w:color"))
#                 if color_el is not None:
#                     val = color_el.get(qn("w:val"), "")
#                     if val.upper() not in ("", "AUTO", "000000"):
#                         self.issues.append(Issue("COLOR-CONTRAST", "WARNING",
#                                                   f"Para: '{para.text[:50]}'",
#                                                   f"Custom font color '{val}' found. Verify high-contrast compliance."))
#                         warned += 1
#         print(f"  [R11] Colors: {warned} custom color runs flagged.")

#     # ── R12: Page / section breaks ───────────────────────────
#     def _check_page_breaks(self):
#         """Warn about unnecessary explicit page breaks (manual paragraph-based breaks)."""
#         warned = 0
#         for para in self.doc.paragraphs:
#             for run in para.runs:
#                 if run._r.find(".//" + qn("w:br")) is not None:
#                     br = run._r.find(".//" + qn("w:br"))
#                     br_type = br.get(qn("w:type"), "")
#                     if br_type in ("page", "column"):
#                         self.issues.append(Issue("PAGEBREAK", "WARNING",
#                                                   f"Para: '{para.text[:50]}'",
#                                                   f"Explicit {br_type} break found. Verify it is necessary."))
#                         warned += 1
#         print(f"  [R12] Page/column breaks: {warned} found.")

#     # ── R13: Bookmarks ───────────────────────────────────────
#     def _check_bookmarks(self):
#         """Warn if bookmarks are present (guidelines say remove default bookmarks)."""
#         bm_starts = []
#         for para in self.doc.paragraphs:
#             for bm in para._p.findall(".//" + qn("w:bookmarkStart")):
#                 bm_name = bm.get(qn("w:name"), "")
#                 if not bm_name.startswith("_"):  # Skip internal Word bookmarks
#                     bm_starts.append(bm_name)

#         if bm_starts:
#             self.issues.append(Issue("BOOKMARKS", "WARNING", "Document",
#                                       f"Found {len(bm_starts)} user bookmark(s): {bm_starts[:5]}. "
#                                       "Remove all bookmarks per guidelines."))
#         print(f"  [R13] Bookmarks: {len(bm_starts)} user bookmarks found.")

#     # ═══════════════════════════════════════════════════════
#     #  SUMMARY REPORT
#     # ═══════════════════════════════════════════════════════
#     def _print_summary(self):
#         errors   = [i for i in self.issues if i.severity == "ERROR"]
#         warnings = [i for i in self.issues if i.severity == "WARNING"]
#         fixed    = [i for i in self.issues if i.fix_applied]

#         print(f"\n{'─'*60}")
#         print(f"  ACCESSIBILITY REPORT SUMMARY")
#         print(f"{'─'*60}")
#         print(f"  Total issues  : {len(self.issues)}")
#         print(f"  ✗ Errors      : {len(errors)}")
#         print(f"  ⚠ Warnings    : {len(warnings)}")
#         print(f"  ✔ Auto-fixed  : {len(fixed)}")
#         print(f"  ✗ Manual fixes: {len(self.issues) - len(fixed)}")
#         print(f"{'─'*60}\n")

#         if errors:
#             print("  ERRORS (must fix):")
#             for i in errors:
#                 status = "✔ FIXED" if i.fix_applied else "✗ MANUAL"
#                 print(f"    [{status}] [{i.rule_id}] {i.location}")
#                 print(f"             {i.description}")

#         if warnings:
#             print("\n  WARNINGS (review recommended):")
#             for i in warnings:
#                 status = "✔ FIXED" if i.fix_applied else "✗ MANUAL"
#                 print(f"    [{status}] [{i.rule_id}] {i.location}")
#                 print(f"             {i.description}")

#         print(f"\n{'='*60}\n")
import re
import sys
from copy import deepcopy
from pathlib import Path
from urllib.parse import urlparse
from docx import Document
from docx.opc.constants import RELATIONSHIP_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Pt


class Issue:
    def __init__(self, rule_id, severity, location, description, fix_applied=False):
        self.rule_id = rule_id
        self.severity = severity
        self.location = location
        self.description = description
        self.fix_applied = fix_applied

    def __repr__(self):
        status = "FIXED" if self.fix_applied else "MANUAL"
        return f"[{self.severity}][{self.rule_id}] {self.location}: {self.description} — {status}"


class AccessibilityChecker:

    HEADING_STYLES = {
        "heading 1": 1, "heading 2": 2, "heading 3": 3,
        "heading 4": 4, "heading 5": 5, "heading 6": 6,
    }
    URL_RE = re.compile(r"(?P<url>(?:https?://|www\.)[^\s<>\"]+)", re.IGNORECASE)
    INTERNAL_REF_RE = re.compile(
        r"\b(?:see|refer(?:ring)? to|as shown in|in)\s+"
        r"(?:the\s+)?(?:section|figure|table|appendix|chapter)\s+[A-Za-z0-9][\w\.\-]*",
        re.IGNORECASE,
    )

    def __init__(
        self,
        input_path,
        output_path=None,
        auto_fix=True,
        preserve_original_formatting=True,
        force_demote_headings=True,
    ):
        self.input_path = Path(input_path)
        self.output_path = Path(output_path) if output_path else self.input_path.with_stem(
            self.input_path.stem + "_accessible"
        )
        self.auto_fix = auto_fix
        self.preserve_original_formatting = preserve_original_formatting
        self.force_demote_headings = force_demote_headings
        self.issues = []
        self.doc = None

    def _allow_formatting_changes(self):
        return self.auto_fix and not self.preserve_original_formatting

    def _allow_text_changes(self):
        """
        Text-mutating fixes must stay off when preserving original formatting,
        to keep generated output word count and textual content identical to input.
        """
        return self.auto_fix and not self.preserve_original_formatting

    # ================= RUN =================
    def run(self):
        print("\n=== Accessibility Checker ===\n")
        print(
            f"Mode: auto_fix={self.auto_fix}, "
            f"preserve_original_formatting={self.preserve_original_formatting}, "
            f"force_demote_headings={self.force_demote_headings}"
        )

        self.doc = Document(str(self.input_path))
        self.issues = []
        header_footer_snapshot = self._snapshot_primary_headers_footers()

        if self.auto_fix:
            # Pass 1: apply automatic fixes.
            self._run_all_checks()
            # Keep original header content exactly as-is.
            # Footer is intentionally not restored because footer normalization
            # is an explicit output requirement.
            self._restore_primary_headers_footers(header_footer_snapshot)
            # Keep only the final post-fix audit results (no duplicate reporting).
            self.issues = []
            # Pass 2: audit after fixes.
            self.auto_fix = False
            self._run_all_checks()
            self.auto_fix = True
        else:
            self._run_all_checks()

        if self.auto_fix:
            final_path = self._save_with_fallback(self.output_path)
            print(f"Saved: {final_path}")

        self._print_summary()
        return self.issues

    def _save_with_fallback(self, preferred_path):
        """
        Save to preferred_path. If the file is locked/open, save to a numbered fallback path.
        """
        try:
            self.doc.save(str(preferred_path))
            return preferred_path
        except PermissionError:
            fallback_path = self._next_available_output_path(preferred_path)
            self.doc.save(str(fallback_path))
            print(f"Output file is in use: {preferred_path}")
            print(f"Saved to fallback file: {fallback_path}")
            return fallback_path

    @staticmethod
    def _next_available_output_path(base_path):
        stem = base_path.stem
        suffix = base_path.suffix or ".docx"
        parent = base_path.parent
        index = 1
        while True:
            candidate = parent / f"{stem}_{index}{suffix}"
            if not candidate.exists():
                return candidate
            index += 1

    def _snapshot_primary_headers_footers(self):
        """
        Preserve primary header XML so output keeps it exactly as input.
        """
        snapshot = []
        for section in self.doc.sections:
            snapshot.append({
                "header": deepcopy(section.header._element),
            })
        return snapshot

    def _restore_primary_headers_footers(self, snapshot):
        for idx, section in enumerate(self.doc.sections):
            if idx >= len(snapshot):
                break
            saved = snapshot[idx]
            self._replace_element_content(section.header._element, saved["header"])

    @staticmethod
    def _replace_element_content(target, source):
        for child in list(target):
            target.remove(child)
        for key in list(target.attrib.keys()):
            del target.attrib[key]
        for key, value in source.attrib.items():
            target.set(key, value)
        target.text = source.text
        target.tail = source.tail
        for child in source:
            target.append(deepcopy(child))

    def _run_all_checks(self):
        self._normalize_headings()
        self._normalize_bookmarks_keep_toc()
        self._normalize_footer_copyright()
        self._check_headings()
        self._check_empty_paragraphs()
        self._check_double_spaces()
        self._check_underlines()
        self._check_lists()
        self._check_tables()
        self._check_links()
        self._check_images()
        self._check_equations()
        self._check_metadata()
        self._check_color_contrast()
        self._check_page_breaks()

    # ================= HEADINGS =================
    def _check_headings(self):
        if self.force_demote_headings:
            return

        headings = []
        h1_blocks = 0
        prev_was_h1 = False

        for para in self.doc.paragraphs:
            style = para.style.name.lower()
            if style in self.HEADING_STYLES:
                level = self.HEADING_STYLES[style]
                headings.append((level, para))
                if level == 1:
                    if not prev_was_h1:
                        h1_blocks += 1
                    prev_was_h1 = True
                else:
                    prev_was_h1 = False
            else:
                prev_was_h1 = False

        if not headings:
            self.issues.append(Issue("NO-HEADINGS", "ERROR", "Doc", "No headings found"))
            return

        first_level, first_para = headings[0]
        if first_level != 1:
            self.issues.append(Issue("H1-NOT-FIRST", "ERROR", first_para.text,
                                     "First heading must be H1"))

        if h1_blocks != 1:
            self.issues.append(Issue("H1-COUNT", "ERROR", "Doc",
                                     f"H1 count = {h1_blocks}"))

    def _normalize_headings(self):
        if not self.auto_fix or not self.force_demote_headings:
            return

        demoted = 0
        for para in self._iter_all_paragraphs():
            if not self._paragraph_has_heading_semantics(para):
                continue

            demoted += 1
            self._freeze_heading_paragraph_indentation(para)
            self._freeze_heading_run_appearance(para)

            if self._style_exists("Normal"):
                para.style = "Normal"

            # Remove XML heading markers to force paragraph back to body text semantics.
            ppr = para._p.get_or_add_pPr()
            pstyle = ppr.find(qn("w:pStyle"))
            if pstyle is not None:
                ppr.remove(pstyle)

            outline_lvl = ppr.find(qn("w:outlineLvl"))
            if outline_lvl is not None:
                ppr.remove(outline_lvl)

        if demoted:
            self.issues.append(Issue(
                "HEADINGS-DEMOTED",
                "WARNING",
                "Doc",
                f"Force-demoted {demoted} heading paragraph(s) to Normal style.",
                fix_applied=True,
            ))

    def _paragraph_has_heading_semantics(self, para):
        style = para.style
        ppr = para._p.find(qn("w:pPr"))
        pstyle_val = ""
        outline_val = None
        if ppr is not None:
            pstyle = ppr.find(qn("w:pStyle"))
            if pstyle is not None:
                pstyle_val = (pstyle.get(qn("w:val")) or "").strip().lower()
            outline_lvl = ppr.find(qn("w:outlineLvl"))
            if outline_lvl is not None:
                outline_val = (outline_lvl.get(qn("w:val")) or "").strip()

        if pstyle_val.startswith("heading"):
            return True
        if outline_val and outline_val.isdigit() and int(outline_val) <= 8:
            return True
        if self._style_chain_has_heading_semantics(style):
            return True
        return False

    def _style_chain_has_heading_semantics(self, style):
        for current in self._iter_style_chain(style):
            if current is None:
                continue

            style_name = (current.name or "").strip().lower()
            style_id = (getattr(current, "style_id", "") or "").strip().lower()
            if style_name in self.HEADING_STYLES:
                return True
            if style_name.startswith("heading "):
                return True
            if style_id.startswith("heading"):
                return True

            style_el = getattr(current, "element", None)
            if style_el is None:
                continue

            ppr = style_el.find(qn("w:pPr"))
            if ppr is None:
                continue

            pstyle = ppr.find(qn("w:pStyle"))
            if pstyle is not None:
                val = (pstyle.get(qn("w:val")) or "").strip().lower()
                if val.startswith("heading"):
                    return True

            outline_lvl = ppr.find(qn("w:outlineLvl"))
            if outline_lvl is not None:
                outline_val = (outline_lvl.get(qn("w:val")) or "").strip()
                if outline_val.isdigit() and int(outline_val) <= 8:
                    return True

        return False

    def _freeze_heading_run_appearance(self, para):
        """
        Preserve heading visual appearance by materializing inherited font/size/color
        into each run before removing heading semantics.
        """
        rfonts_el = self._resolve_style_rpr_property(para.style, qn("w:rFonts"))
        color_el = self._resolve_style_rpr_property(para.style, qn("w:color"))
        size_el = self._resolve_style_rpr_property(para.style, qn("w:sz"))
        size_cs_el = self._resolve_style_rpr_property(para.style, qn("w:szCs"))

        for run in para.runs:
            rpr = run._r.find(qn("w:rPr"))
            if rpr is None:
                rpr = OxmlElement("w:rPr")
                run._r.insert(0, rpr)

            if rfonts_el is not None and rpr.find(qn("w:rFonts")) is None:
                rpr.append(deepcopy(rfonts_el))
            if color_el is not None and rpr.find(qn("w:color")) is None:
                rpr.append(deepcopy(color_el))
            if size_el is not None and rpr.find(qn("w:sz")) is None:
                rpr.append(deepcopy(size_el))
            if size_cs_el is not None and rpr.find(qn("w:szCs")) is None:
                rpr.append(deepcopy(size_cs_el))

    def _freeze_heading_paragraph_indentation(self, para):
        """
        Preserve effective paragraph indentation/alignment before demoting heading style.
        """
        ppr = para._p.get_or_add_pPr()

        if ppr.find(qn("w:jc")) is None:
            jc_el = self._resolve_style_ppr_property(para.style, qn("w:jc"))
            if jc_el is not None:
                ppr.append(deepcopy(jc_el))

        if ppr.find(qn("w:ind")) is None:
            ind_el = self._resolve_style_ppr_property(para.style, qn("w:ind"))
            if ind_el is not None:
                ppr.append(deepcopy(ind_el))

        # Tabs often participate in visual indentation for heading lines.
        if ppr.find(qn("w:tabs")) is None:
            tabs_el = self._resolve_style_ppr_property(para.style, qn("w:tabs"))
            if tabs_el is not None:
                ppr.append(deepcopy(tabs_el))

    def _resolve_style_rpr_property(self, style, tag_name):
        for current in self._iter_style_chain(style):
            if current is None:
                continue
            style_el = getattr(current, "element", None)
            if style_el is None:
                continue
            rpr = style_el.find(qn("w:rPr"))
            if rpr is None:
                continue
            prop = rpr.find(tag_name)
            if prop is not None:
                return prop
        return None

    def _resolve_style_ppr_property(self, style, tag_name):
        for current in self._iter_style_chain(style):
            if current is None:
                continue
            style_el = getattr(current, "element", None)
            if style_el is None:
                continue
            ppr = style_el.find(qn("w:pPr"))
            if ppr is None:
                continue
            prop = ppr.find(tag_name)
            if prop is not None:
                return prop
        return None

    # ================= BOOKMARKS =================
    def _normalize_bookmarks_keep_toc(self):
        removed = self._remove_non_toc_bookmarks() if self.auto_fix else 0
        if removed:
            self.issues.append(Issue(
                "BOOKMARKS-REMOVED",
                "WARNING",
                "Doc",
                f"Removed {removed} non-TOC bookmark(s).",
                fix_applied=True,
            ))

    def _remove_non_toc_bookmarks(self):
        roots = self._bookmark_roots()
        remove_ids = set()
        removed_starts = 0

        for root in roots:
            starts = list(root.findall(".//" + qn("w:bookmarkStart")))
            for bm in starts:
                name = (bm.get(qn("w:name")) or "").strip()
                if self._is_toc_bookmark_name(name):
                    continue
                bm_id = bm.get(qn("w:id"))
                if bm_id is not None:
                    remove_ids.add(bm_id)
                parent = bm.getparent()
                if parent is not None:
                    parent.remove(bm)
                    removed_starts += 1

        removed_ends = 0
        if remove_ids:
            for root in roots:
                ends = list(root.findall(".//" + qn("w:bookmarkEnd")))
                for bm_end in ends:
                    bm_id = bm_end.get(qn("w:id"))
                    if bm_id not in remove_ids:
                        continue
                    parent = bm_end.getparent()
                    if parent is not None:
                        parent.remove(bm_end)
                        removed_ends += 1

        return removed_starts + removed_ends

    def _bookmark_roots(self):
        roots = [self.doc.part.element]
        seen = {id(self.doc.part.element)}

        for footer in self._iter_all_footers():
            root = footer._element
            if id(root) in seen:
                continue
            roots.append(root)
            seen.add(id(root))

        for section in self.doc.sections:
            header = section.header
            if header is None:
                continue
            root = header._element
            if id(root) in seen:
                continue
            roots.append(root)
            seen.add(id(root))

        return roots

    @staticmethod
    def _is_toc_bookmark_name(name):
        lower = (name or "").strip().lower()
        return lower.startswith("_toc") or lower.startswith("toc")

    # ================= FOOTERS =================
    def _normalize_footer_copyright(self):
        if not self.auto_fix or self.preserve_original_formatting:
            return

        updated_runs = 0
        for footer in self._iter_all_footers():
            for para in self._iter_story_paragraphs(footer):
                if not self._is_copyright_paragraph(para):
                    continue

                source_font = self._resolve_source_font_name(para)
                if self._normalize_paragraph_sentence_case(para):
                    updated_runs += len(para.runs)

                for run in para.runs:
                    if self._apply_run_font_size_and_style(run, source_font):
                        updated_runs += 1

        if updated_runs:
            self.issues.append(Issue(
                "COPYRIGHT-STYLE",
                "WARNING",
                "Footer",
                (
                    f"Normalized {updated_runs} copyright footer run(s): "
                    "source font, 10 pt, sentence case, non-italic."
                ),
                fix_applied=True,
            ))

    def _iter_all_footers(self):
        seen = set()
        for section in self.doc.sections:
            candidates = [section.footer]
            first_page_footer = getattr(section, "first_page_footer", None)
            even_page_footer = getattr(section, "even_page_footer", None)
            if first_page_footer is not None:
                candidates.append(first_page_footer)
            if even_page_footer is not None:
                candidates.append(even_page_footer)

            for footer in candidates:
                if footer is None:
                    continue
                key = id(footer._element)
                if key in seen:
                    continue
                seen.add(key)
                yield footer

    def _iter_story_paragraphs(self, story):
        for para in story.paragraphs:
            yield para
        for table in self._iter_tables(story.tables):
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        yield para

    @staticmethod
    def _is_copyright_paragraph(para):
        text = (para.text or "").strip().lower()
        return bool(text) and ("copyright" in text or "©" in text or "(c)" in text)

    @staticmethod
    def _to_sentence_case(text):
        lowered = (text or "").lower()
        chars = list(lowered)
        capitalize_next = True

        for idx, ch in enumerate(chars):
            if capitalize_next and ch.isalpha():
                chars[idx] = ch.upper()
                capitalize_next = False
            if ch in ".!?":
                capitalize_next = True

        return "".join(chars)

    def _normalize_paragraph_sentence_case(self, para):
        runs = list(para.runs)
        if not runs:
            return False

        original = "".join(run.text or "" for run in runs)
        converted = self._to_sentence_case(original)
        if converted == original:
            return False

        index = 0
        for run in runs:
            length = len(run.text or "")
            if length == 0:
                continue
            run.text = converted[index:index + length]
            index += length
        return True

    @staticmethod
    def _resolve_source_font_name(para):
        for run in para.runs:
            if run.font.name:
                return run.font.name
        return None

    def _apply_run_font_size_and_style(self, run, source_font):
        changed = False

        if source_font and run.font.name != source_font:
            run.font.name = source_font
            changed = True
        if run.font.size is None or abs(float(run.font.size.pt) - 10.0) > 0.01:
            run.font.size = Pt(10)
            changed = True
        if run.italic:
            run.italic = False
            changed = True

        rpr = run._r.find(qn("w:rPr"))
        if rpr is None:
            rpr = OxmlElement("w:rPr")
            run._r.insert(0, rpr)

        if source_font:
            rfonts = rpr.find(qn("w:rFonts"))
            if rfonts is None:
                rfonts = OxmlElement("w:rFonts")
                rpr.append(rfonts)

            for key in ("ascii", "hAnsi", "cs", "eastAsia"):
                attr = qn(f"w:{key}")
                if rfonts.get(attr) != source_font:
                    rfonts.set(attr, source_font)
                    changed = True

        return changed

    # ================= PARAGRAPHS =================
    def _check_empty_paragraphs(self):
        # Empty-paragraph cleanup is allowed in preserve mode because it does not
        # alter textual content/word count; spacing is transferred to neighbors.
        if self._allow_text_changes():
            self._replace_double_line_breaks_with_paragraphs(self.doc)
            for cell in self._iter_table_cells():
                self._replace_double_line_breaks_with_paragraphs(cell)

        self._normalize_empty_paragraphs(self.doc)
        for cell in self._iter_table_cells():
            self._normalize_empty_paragraphs(cell)

    def _iter_table_cells(self):
        """
        Yield table cells from all tables, including nested tables.
        Do not de-duplicate by object id; merged-table proxies can cause misses.
        """
        for table in self._iter_tables(self.doc.tables):
            for row in table.rows:
                for cell in row.cells:
                    yield cell

    def _iter_tables(self, tables):
        for table in tables:
            yield table
            for row in table.rows:
                for cell in row.cells:
                    if cell.tables:
                        yield from self._iter_tables(cell.tables)

    def _normalize_empty_paragraphs(self, container):
        """
        Remove empty paragraph blocks and preserve visual spacing by transferring
        equivalent gap to a nearby non-empty paragraph.
        """
        i = 0
        paragraphs = container.paragraphs
        while i < len(paragraphs):
            if not self._is_visually_empty_paragraph(paragraphs[i]):
                i += 1
                continue

            j = i
            while (
                j < len(paragraphs)
                and self._is_visually_empty_paragraph(paragraphs[j])
                and (j == i or self._paragraphs_are_xml_neighbors(paragraphs[j - 1], paragraphs[j]))
            ):
                j += 1

            if not self.auto_fix:
                i = j
                continue

            empty_block = paragraphs[i:j]
            removable = [p for p in empty_block if self._can_remove_paragraph_element(p)]
            if removable:
                anchor_para, use_before = self._select_anchor_for_empty_block(paragraphs, i, j, empty_block)
                if anchor_para is not None:
                    added_space_pt = self._estimate_empty_gap_pt(removable)
                    self._transfer_spacing_to_anchor(anchor_para, added_space_pt, use_before)

            for para in removable:
                para._element.getparent().remove(para._element)

            paragraphs = container.paragraphs
            if not removable:
                # Some containers (e.g., a lone table-cell paragraph) require at least
                # one paragraph node; if nothing was removed, advance to avoid stalling.
                i = j

    def _replace_double_line_breaks_with_paragraphs(self, container):
        """
        Replace repeated manual line breaks with paragraph splits.
        Equivalent to repeated replacement: ^l^l -> ^p
        """
        if not self.auto_fix:
            return

        idx = 0
        while idx < len(container.paragraphs):
            para = container.paragraphs[idx]
            text = (para.text or "").replace("\r", "")
            if "\n\n" not in text:
                idx += 1
                continue

            parts = self._split_on_double_line_breaks(text)
            if len(parts) <= 1:
                idx += 1
                continue

            self._replace_paragraph_with_parts(container, para, parts)
            idx += len(parts)

    @staticmethod
    def _split_on_double_line_breaks(text):
        parts = []
        buf = []
        i = 0
        while i < len(text):
            if i + 1 < len(text) and text[i] == "\n" and text[i + 1] == "\n":
                parts.append("".join(buf))
                buf = []
                i += 2
                continue
            buf.append(text[i])
            i += 1
        parts.append("".join(buf))
        return parts

    def _replace_paragraph_with_parts(self, container, paragraph, parts):
        self._set_paragraph_text_with_breaks(paragraph, parts[0])
        anchor = paragraph
        for part_text in parts[1:]:
            new_para = container.add_paragraph()
            self._copy_paragraph_properties(paragraph, new_para)
            self._set_paragraph_text_with_breaks(new_para, part_text)
            anchor._element.addnext(new_para._element)
            anchor = new_para

    @staticmethod
    def _set_paragraph_text_with_breaks(para, text):
        p_el = para._p
        for child in list(p_el):
            if child.tag == qn("w:pPr"):
                continue
            p_el.remove(child)
        if text:
            para.add_run(text)

    @staticmethod
    def _copy_paragraph_properties(source_para, target_para):
        """
        Copy paragraph-level properties (alignment, indentation, tabs, spacing, etc.)
        so split paragraphs preserve original layout semantics.
        """
        src_ppr = source_para._p.find(qn("w:pPr"))
        tgt_p = target_para._p
        tgt_ppr = tgt_p.find(qn("w:pPr"))
        if tgt_ppr is not None:
            tgt_p.remove(tgt_ppr)

        if src_ppr is not None:
            tgt_p.insert(0, deepcopy(src_ppr))
            return

        try:
            target_para.style = source_para.style
        except Exception:
            pass

    def _can_remove_paragraph_element(self, para):
        parent = para._element.getparent()
        if parent is None:
            return False

        required_min_paragraphs = 1 if parent.tag in {qn("w:tc"), qn("w:txbxContent")} else 0
        total_paragraphs = sum(1 for child in parent if child.tag == qn("w:p"))
        return total_paragraphs > required_min_paragraphs

    def _paragraphs_are_xml_neighbors(self, prev_para, next_para):
        if prev_para is None or next_para is None:
            return False
        return prev_para._element.getnext() is next_para._element

    def _select_anchor_for_empty_block(self, paragraphs, start, end, empty_block):
        prev_idx = start - 1
        while prev_idx >= 0 and self._is_visually_empty_paragraph(paragraphs[prev_idx]):
            prev_idx -= 1

        next_idx = end
        while next_idx < len(paragraphs) and self._is_visually_empty_paragraph(paragraphs[next_idx]):
            next_idx += 1

        prev_para = paragraphs[prev_idx] if prev_idx >= 0 else None
        next_para = paragraphs[next_idx] if next_idx < len(paragraphs) else None
        if prev_para is None and next_para is None:
            return None, False

        left_neighbor = self._adjacent_non_property_sibling(empty_block[0]._element, direction=-1)
        right_neighbor = self._adjacent_non_property_sibling(empty_block[-1]._element, direction=1)
        left_is_table = left_neighbor is not None and left_neighbor.tag == qn("w:tbl")
        right_is_table = right_neighbor is not None and right_neighbor.tag == qn("w:tbl")

        # Priority: between table and paragraph, always anchor to paragraph side.
        if left_is_table and next_para is not None:
            return next_para, True
        if right_is_table and prev_para is not None:
            return prev_para, False

        if prev_para is None:
            return next_para, True
        if next_para is None:
            return prev_para, False

        dist_prev = start - prev_idx
        dist_next = next_idx - (end - 1)
        if dist_next < dist_prev:
            return next_para, True
        return prev_para, False

    def _adjacent_non_property_sibling(self, element, direction):
        current = element
        while True:
            current = current.getprevious() if direction < 0 else current.getnext()
            if current is None:
                return None
            if current.tag in {qn("w:tcPr"), qn("w:trPr"), qn("w:tblPr"), qn("w:pPr")}:
                continue
            return current

    def _compute_removable_empty_count(self, empty_block):
        if not empty_block:
            return 0

        parent = empty_block[0]._element.getparent()
        if parent is None:
            return 0

        # Keep one paragraph in cell/textbox containers to avoid corruption.
        required_min_paragraphs = 1 if parent.tag in {qn("w:tc"), qn("w:txbxContent")} else 0
        total_paragraphs = 0
        for child in parent:
            if child.tag == qn("w:p"):
                total_paragraphs += 1

        max_removable = max(0, total_paragraphs - required_min_paragraphs)
        return min(len(empty_block), max_removable)

    def _is_visually_empty_paragraph(self, para):
        raw_text = para.text or ""

        text = raw_text.replace("\u200b", "").replace("\u00a0", " ").strip()
        if text:
            return False

        # Keep paragraphs that carry visible content via objects/drawings.
        p_el = para._p
        if p_el.find(".//" + qn("w:drawing")) is not None:
            return False
        if p_el.find(".//" + qn("w:object")) is not None:
            return False
        if p_el.find(".//" + qn("w:pict")) is not None:
            return False
        # Preserve page/section structure to keep pagination intact.
        if p_el.find(".//" + qn("w:br")) is not None:
            return False
        if p_el.find(".//" + qn("w:lastRenderedPageBreak")) is not None:
            return False
        if p_el.find(".//" + qn("w:sectPr")) is not None:
            return False
        return True

    def _estimate_empty_gap_pt(self, empty_paras):
        if not empty_paras:
            return 0.0

        total = 0.0
        for para in empty_paras:
            para_gap = self._estimate_single_empty_para_gap_pt(para)
            total += para_gap
        return total

    def _estimate_single_empty_para_gap_pt(self, para):
        line_pt = self._resolve_effective_line_height_pt(para)
        before = self._resolve_effective_space_pt(para, kind="before")
        after = self._resolve_effective_space_pt(para, kind="after")
        return max(8.0, line_pt) + before + after

    def _resolve_base_font_size_pt(self, para):
        for run in para.runs:
            if run.font.size is not None:
                return float(run.font.size.pt)
        for style in self._iter_style_chain(para.style):
            if style.font is not None and style.font.size is not None:
                return float(style.font.size.pt)
        return 12.0

    def _iter_style_chain(self, style):
        seen = set()
        current = style
        while current is not None and id(current) not in seen:
            seen.add(id(current))
            yield current
            current = current.base_style

    def _resolve_effective_space_pt(self, para, kind):
        if kind == "before":
            direct = para.paragraph_format.space_before
        else:
            direct = para.paragraph_format.space_after
        if direct is not None:
            return float(direct.pt)

        for style in self._iter_style_chain(para.style):
            pf = style.paragraph_format
            candidate = pf.space_before if kind == "before" else pf.space_after
            if candidate is not None:
                return float(candidate.pt)
        return 0.0

    def _resolve_effective_line_height_pt(self, para):
        line_spacing = para.paragraph_format.line_spacing
        if line_spacing is None:
            for style in self._iter_style_chain(para.style):
                line_spacing = style.paragraph_format.line_spacing
                if line_spacing is not None:
                    break

        if hasattr(line_spacing, "pt"):
            return float(line_spacing.pt)
        if isinstance(line_spacing, (int, float)):
            return float(line_spacing) * self._resolve_base_font_size_pt(para)
        return self._resolve_base_font_size_pt(para)

    def _transfer_spacing_to_anchor(self, anchor_para, total_height_pt, use_before):
        if anchor_para is None or total_height_pt <= 0:
            return

        extra = Pt(total_height_pt)
        if use_before:
            self._add_space_before(anchor_para, extra)
        else:
            self._add_space_after(anchor_para, extra)

    def _add_space_after(self, para, extra):
        current = self._resolve_effective_space_pt(para, kind="after")
        para.paragraph_format.space_after = Pt(current + extra.pt)

    def _add_space_before(self, para, extra):
        current = self._resolve_effective_space_pt(para, kind="before")
        para.paragraph_format.space_before = Pt(current + extra.pt)

    def _check_double_spaces(self):
        fixed_runs = 0
        for para in self._iter_all_paragraphs():
            for run in para.runs:
                text = run.text or ""
                if "  " not in text:
                    continue

                normalized = self._normalize_spacing_runs(text)
                changed = normalized != text
                if not changed:
                    continue

                if self.auto_fix:
                    run.text = normalized
                fixed_runs += 1
                self.issues.append(Issue(
                    "DOUBLE-SPACE",
                    "WARNING",
                    para.text[:500],
                    "Multiple spaces used for indenting/alignment were normalized.",
                    fix_applied=self.auto_fix,
                ))

        return fixed_runs

    @staticmethod
    def _normalize_spacing_runs(text):
        """
        Normalize multiple spaces:
        - Collapse leading indentation spaces to a single space.
        - Collapse any 3+ spaces to one.
        - Collapse remaining double spaces to one, except likely intentional
          sentence spacing after punctuation (e.g., '.  Next').
        """
        if not text:
            return text

        normalized = re.sub(r"(^|\n) {2,}", r"\1 ", text)
        normalized = re.sub(r" {3,}", " ", normalized)

        intentional_sentence_gap = re.compile(
            r"(?<=[.!?])  (?=(?:[\"'\(\[])?[A-Z0-9])"
        )

        protected = "\uFFF0"
        normalized = intentional_sentence_gap.sub(protected, normalized)
        normalized = normalized.replace("  ", " ")
        normalized = normalized.replace(protected, "  ")
        return normalized

    def _check_underlines(self):
        for para in self._iter_all_paragraphs():
            teaching_context = self._paragraph_mentions_underlining_concept(para.text)
            for run in para.runs:
                if not run.underline:
                    continue
                if self._run_is_in_hyperlink(run):
                    continue
                if teaching_context:
                    continue

                fixed = False
                if self.auto_fix:
                    run.underline = False
                    fixed = True

                self.issues.append(Issue(
                    "UNDERLINE-NONLINK",
                    "ERROR",
                    para.text[:500],
                    "Underline used on non-link text.",
                    fix_applied=fixed,
                ))

    def _check_lists(self):
        for para in self._iter_all_paragraphs():
            info = self._ordered_list_info(para.text)
            if info is None:
                continue
            if self._paragraph_is_numbered_list(para):
                continue

            fixed = False
            if self._allow_formatting_changes():
                fixed = self._apply_ordered_list_format(para, info)

            self.issues.append(Issue(
                "LIST",
                "ERROR",
                para.text[:500],
                "Manual ordered list detected",
                fix_applied=fixed,
            ))

    def _iter_all_paragraphs(self):
        for para in self.doc.paragraphs:
            yield para
        for cell in self._iter_table_cells():
            for para in cell.paragraphs:
                yield para

    def _ordered_list_info(self, text):
        raw = text or ""
        if not raw.strip():
            return None

        # Hierarchical numeric: 1.2 / 1.2.3 / 2.1)
        m = re.match(r"^(\s*)(\d+(?:\.\d+)+)[\.\)]?\s+(.*)$", raw)
        if m:
            indent = len(m.group(1).expandtabs(4)) // 4
            level = min(3, max(1, m.group(2).count(".") + 1 + indent))
            return {"level": level, "content": m.group(3).strip(), "prefix_len": m.start(3)}

        # Numeric: 1. text / 1) text
        m = re.match(r"^(\s*)(\d+)[\.\)]\s+(.*)$", raw)
        if m:
            indent = len(m.group(1).expandtabs(4)) // 4
            level = min(3, max(1, 1 + indent))
            return {"level": level, "content": m.group(3).strip(), "prefix_len": m.start(3)}

        # Alphabetic: a) text / A. text
        m = re.match(r"^(\s*)([A-Za-z])[\.\)]\s+(.*)$", raw)
        if m:
            indent = len(m.group(1).expandtabs(4)) // 4
            level = min(3, max(1, 2 + indent))
            return {"level": level, "content": m.group(3).strip(), "prefix_len": m.start(3)}

        # Roman numeral: i) text / IV. text
        m = re.match(r"^(\s*)([ivxlcdmIVXLCDM]+)[\.\)]\s+(.*)$", raw)
        if m:
            indent = len(m.group(1).expandtabs(4)) // 4
            level = min(3, max(1, 2 + indent))
            return {"level": level, "content": m.group(3).strip(), "prefix_len": m.start(3)}

        return None

    def _paragraph_is_numbered_list(self, para):
        style_name = (para.style.name or "").lower()
        if style_name.startswith("list number"):
            return True

        ppr = para._p.find(qn("w:pPr"))
        if ppr is None:
            return False
        num_pr = ppr.find(qn("w:numPr"))
        return num_pr is not None

    def _style_exists(self, style_name):
        try:
            self.doc.styles[style_name]
            return True
        except KeyError:
            return False

    def _apply_ordered_list_format(self, para, info):
        level = int(info["level"])
        preferred = {1: "List Number", 2: "List Number 2", 3: "List Number 3"}
        style_name = preferred.get(level, "List Number")
        if not self._style_exists(style_name):
            if self._style_exists("List Number"):
                style_name = "List Number"
            else:
                # Do not alter text if we cannot apply a true numbered-list style.
                return False

        para.style = style_name
        self._strip_manual_list_prefix_preserve_runs(para, int(info.get("prefix_len", 0)))
        return True

    def _strip_manual_list_prefix_preserve_runs(self, para, prefix_len):
        """
        Remove leading manual marker text from paragraph runs while preserving run formatting.
        """
        if prefix_len <= 0:
            return

        runs = list(para.runs)
        if not runs:
            return

        remaining = prefix_len
        for run in runs:
            if remaining <= 0:
                break
            text = run.text or ""
            if not text:
                continue
            take = min(len(text), remaining)
            run.text = text[take:]
            remaining -= take

    # ================= TABLES =================
    def _check_tables(self):
        all_tables = list(self._iter_tables(self.doc.tables))
        for idx, table in enumerate(all_tables, start=1):
            location = f"Table {idx}"

            if self._is_probable_layout_table(table):
                self.issues.append(Issue(
                    "TABLE-LAYOUT",
                    "WARNING",
                    location,
                    "Table appears to be used for layout rather than data.",
                    fix_applied=False,
                ))

            if self.auto_fix and self._allow_formatting_changes():
                self._normalize_table_design_options(table)
                self._disable_row_break_across_pages(table)

            # Remove blank rows
            for row in list(table.rows):
                if all(cell.text.strip() == "" for cell in row.cells):
                    fixed = False
                    if self._allow_formatting_changes():
                        row._tr.getparent().remove(row._tr)
                        fixed = True
                    self.issues.append(Issue(
                        "TABLE-BLANK-ROW",
                        "ERROR",
                        location,
                        "Completely blank row found in table.",
                        fix_applied=fixed,
                    ))

            # Report blank columns
            row_count = len(table.rows)
            if row_count > 0:
                max_cols = max((len(r.cells) for r in table.rows), default=0)
                for c_idx in range(max_cols):
                    col_blank = True
                    for r_idx in range(row_count):
                        if c_idx >= len(table.rows[r_idx].cells):
                            continue
                        if table.rows[r_idx].cells[c_idx].text.strip() != "":
                            col_blank = False
                            break
                    if col_blank:
                        self.issues.append(Issue(
                            "TABLE-BLANK-COL",
                            "ERROR",
                            location,
                            f"Completely blank column found in table (column {c_idx + 1}).",
                            fix_applied=False,
                        ))

            # Header row
            if table.rows:
                tr = table.rows[0]._tr
                trPr = tr.find(qn("w:trPr"))
                if trPr is None:
                    trPr = OxmlElement("w:trPr")
                    tr.insert(0, trPr)

                if trPr.find(qn("w:tblHeader")) is None:
                    if self.auto_fix:
                        trPr.append(OxmlElement("w:tblHeader"))

    @staticmethod
    def _is_probable_layout_table(table):
        """
        Heuristic for layout-only tables (non-data usage).
        """
        rows = len(table.rows)
        cols = len(table.columns) if table.columns else 0
        if rows <= 1 or cols <= 1:
            return True

        total = max(1, rows * max(1, cols))
        non_empty = 0
        very_long = 0
        for row in table.rows:
            for cell in row.cells:
                txt = (cell.text or "").strip()
                if txt:
                    non_empty += 1
                if len(txt) > 160:
                    very_long += 1

        fill_ratio = non_empty / total
        if fill_ratio < 0.35:
            return True
        if very_long > 0 and non_empty <= max(2, rows):
            return True
        return False

    def _normalize_table_design_options(self, table):
        """
        Enforce table style options:
        - Header Row: checked (w:firstRow=1)
        - First Column: checked (w:firstColumn=1)
        - Banded Rows: checked (w:noHBand=0)
        - Banded Columns: checked (w:noVBand=0)
        - Total Row: unchecked (w:lastRow=0)
        - Last Column: unchecked (w:lastColumn=0)
        """
        tbl_pr = table._tbl.tblPr
        if tbl_pr is None:
            tbl_pr = OxmlElement("w:tblPr")
            table._tbl.insert(0, tbl_pr)

        tbl_look = tbl_pr.find(qn("w:tblLook"))
        if tbl_look is None:
            tbl_look = OxmlElement("w:tblLook")
            tbl_pr.append(tbl_look)

        # Explicitly enforce the requested Table Design options.
        tbl_look.set(qn("w:firstRow"), "1")
        tbl_look.set(qn("w:firstColumn"), "1")
        tbl_look.set(qn("w:noHBand"), "0")
        tbl_look.set(qn("w:noVBand"), "0")
        tbl_look.set(qn("w:lastRow"), "0")
        tbl_look.set(qn("w:lastColumn"), "0")

        # Remove cached bitmask value so explicit flags above are authoritative.
        if qn("w:val") in tbl_look.attrib:
            del tbl_look.attrib[qn("w:val")]

    def _disable_row_break_across_pages(self, table):
        """
        Enforce Row -> "Allow row to break across pages" unchecked for all rows.
        WordprocessingML: w:cantSplit present means do not split row across pages.
        """
        for row in table.rows:
            tr = row._tr
            tr_pr = tr.find(qn("w:trPr"))
            if tr_pr is None:
                tr_pr = OxmlElement("w:trPr")
                tr.insert(0, tr_pr)

            if tr_pr.find(qn("w:cantSplit")) is None:
                tr_pr.append(OxmlElement("w:cantSplit"))

    # ================= LINKS =================
    def _check_links(self):
        bookmark_names = self._collect_bookmark_names()
        for para in self._iter_all_paragraphs():
            autolinked = self._autolink_plain_urls_in_paragraph(para) if self._allow_text_changes() else 0
            if autolinked:
                self.issues.append(Issue(
                    "LINK-AUTOFIX",
                    "WARNING",
                    para.text[:500],
                    f"Converted {autolinked} plain URL(s) to clickable hyperlink(s).",
                    fix_applied=True,
                ))

            hyperlink_elements = para._p.findall(".//" + qn("w:hyperlink"))
            has_hyperlink = bool(hyperlink_elements)

            if self._looks_like_internal_reference(para.text) and not has_hyperlink:
                self.issues.append(Issue(
                    "CROSSREF-NOT-LINK",
                    "WARNING",
                    para.text[:500],
                    "Reference to this document appears to be plain text; format it as an internal hyperlink.",
                    fix_applied=False,
                ))

            for hl in hyperlink_elements:
                text = self._hyperlink_text(hl)
                r_id = hl.get(qn("r:id"))
                anchor = hl.get(qn("w:anchor"))

                if r_id:
                    target = self._resolve_external_hyperlink_target(para.part, r_id)
                    if not target:
                        self.issues.append(Issue(
                            "LINK-EXTERNAL-BROKEN",
                            "ERROR",
                            text or para.text[:500],
                            "External hyperlink target is missing/broken.",
                            fix_applied=False,
                        ))
                    elif not self._is_valid_external_target(target):
                        self.issues.append(Issue(
                            "LINK-EXTERNAL-INVALID",
                            "ERROR",
                            text or para.text[:500],
                            f"External hyperlink target is invalid: {target}",
                            fix_applied=False,
                        ))
                elif anchor:
                    if anchor not in bookmark_names:
                        self.issues.append(Issue(
                            "LINK-INTERNAL-BROKEN",
                            "ERROR",
                            text or para.text[:500],
                            f"Internal link anchor '{anchor}' does not exist.",
                            fix_applied=False,
                        ))
                else:
                    self.issues.append(Issue(
                        "LINK-MALFORMED",
                        "ERROR",
                        text or para.text[:500],
                        "Hyperlink has neither external target nor internal anchor.",
                        fix_applied=False,
                    ))

                missing_underline = self._hyperlink_missing_underline(hl)
                fixed = False
                if missing_underline and self.auto_fix:
                    self._force_hyperlink_underline(hl)
                    fixed = True
                if missing_underline:
                    self.issues.append(Issue(
                        "LINK-UNDERLINE",
                        "ERROR",
                        text or para.text[:500],
                        "Hyperlink is not underlined.",
                        fix_applied=fixed,
                    ))

    @staticmethod
    def _paragraph_mentions_underlining_concept(text):
        sample = (text or "").lower()
        return any(token in sample for token in ("underline", "underlined", "underlining"))

    @staticmethod
    def _run_is_in_hyperlink(run):
        parent = run._r.getparent()
        return parent is not None and parent.tag == qn("w:hyperlink")

    @staticmethod
    def _looks_like_internal_reference(text):
        return bool(AccessibilityChecker.INTERNAL_REF_RE.search(text or ""))

    @staticmethod
    def _hyperlink_text(hyperlink_el):
        parts = []
        for t in hyperlink_el.findall(".//" + qn("w:t")):
            parts.append(t.text or "")
        return "".join(parts).strip()

    def _collect_bookmark_names(self):
        names = set()
        for root in self._bookmark_roots():
            starts = root.findall(".//" + qn("w:bookmarkStart"))
            for bm in starts:
                name = (bm.get(qn("w:name")) or "").strip()
                if name:
                    names.add(name)
        return names

    @staticmethod
    def _resolve_external_hyperlink_target(part, r_id):
        try:
            rel = part.rels[r_id]
        except KeyError:
            return None
        if rel.reltype != RELATIONSHIP_TYPE.HYPERLINK:
            return None
        target = getattr(rel, "target_ref", None)
        return (target or "").strip()

    @staticmethod
    def _is_valid_external_target(target):
        parsed = urlparse(target or "")
        if parsed.scheme in {"http", "https"}:
            return bool(parsed.netloc)
        if parsed.scheme == "mailto":
            return "@" in parsed.path
        return False

    @staticmethod
    def _hyperlink_missing_underline(hyperlink_el):
        for run_el in hyperlink_el.findall(".//" + qn("w:r")):
            rpr = run_el.find(qn("w:rPr"))
            if rpr is None:
                return True
            u_el = rpr.find(qn("w:u"))
            if u_el is None:
                return True
            if u_el.get(qn("w:val")) == "none":
                return True
        return False

    @staticmethod
    def _force_hyperlink_underline(hyperlink_el):
        for run_el in hyperlink_el.findall(".//" + qn("w:r")):
            rpr = run_el.find(qn("w:rPr"))
            if rpr is None:
                rpr = OxmlElement("w:rPr")
                run_el.insert(0, rpr)
            u_el = rpr.find(qn("w:u"))
            if u_el is None:
                u_el = OxmlElement("w:u")
                rpr.append(u_el)
            u_el.set(qn("w:val"), "single")

    def _autolink_plain_urls_in_paragraph(self, para):
        """
        Convert plain http(s)/www URLs in non-hyperlink runs to clickable hyperlinks.
        """
        inserted = 0
        runs = list(para.runs)
        for run in runs:
            if self._run_is_in_hyperlink(run):
                continue
            text = run.text or ""
            if not text.strip():
                continue
            if not self.URL_RE.search(text):
                continue

            segments = []
            cursor = 0
            for match in self.URL_RE.finditer(text):
                start, end = match.span("url")
                if start > cursor:
                    segments.append(("text", text[cursor:start], None))
                token = match.group("url")
                core, suffix = self._split_trailing_url_punctuation(token)
                if core:
                    target = self._normalize_url_target(core)
                    segments.append(("link", core, target))
                if suffix:
                    segments.append(("text", suffix, None))
                cursor = end
            if cursor < len(text):
                segments.append(("text", text[cursor:], None))

            if not any(kind == "link" for kind, _, _ in segments):
                continue

            run_el = run._r
            parent = run_el.getparent()
            if parent is None:
                continue
            idx = parent.index(run_el)

            for kind, seg_text, target in segments:
                if not seg_text:
                    continue
                if kind == "text":
                    parent.insert(idx, self._clone_run_with_text(run_el, seg_text))
                    idx += 1
                    continue

                rel_id = para.part.relate_to(target, RELATIONSHIP_TYPE.HYPERLINK, is_external=True)
                hyperlink_el = self._build_hyperlink_element(run_el, seg_text, rel_id)
                parent.insert(idx, hyperlink_el)
                idx += 1
                inserted += 1

            parent.remove(run_el)
        return inserted

    @staticmethod
    def _split_trailing_url_punctuation(token):
        trailing = ""
        core = token or ""
        while core and core[-1] in ".,;:!?)]":
            trailing = core[-1] + trailing
            core = core[:-1]
        return core, trailing

    @staticmethod
    def _normalize_url_target(url_text):
        if (url_text or "").lower().startswith("www."):
            return "https://" + url_text
        return url_text

    @staticmethod
    def _clone_run_with_text(template_run_el, text):
        run_el = OxmlElement("w:r")
        rpr = template_run_el.find(qn("w:rPr"))
        if rpr is not None:
            run_el.append(deepcopy(rpr))
        t_el = OxmlElement("w:t")
        if text.startswith(" ") or text.endswith(" "):
            t_el.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
        t_el.text = text
        run_el.append(t_el)
        return run_el

    def _build_hyperlink_element(self, template_run_el, text, rel_id):
        hl = OxmlElement("w:hyperlink")
        hl.set(qn("r:id"), rel_id)

        run_el = self._clone_run_with_text(template_run_el, text)
        rpr = run_el.find(qn("w:rPr"))
        if rpr is None:
            rpr = OxmlElement("w:rPr")
            run_el.insert(0, rpr)

        u_el = rpr.find(qn("w:u"))
        if u_el is None:
            u_el = OxmlElement("w:u")
            rpr.append(u_el)
        u_el.set(qn("w:val"), "single")

        hl.append(run_el)
        return hl

    # ================= IMAGES =================
    def _check_images(self):
        for para in self._iter_all_paragraphs():
            drawings = para._p.findall(".//" + qn("wp:inline")) + para._p.findall(".//" + qn("wp:anchor"))
            for drawing in drawings:
                docPr = drawing.find(qn("wp:docPr"))
                if docPr is None or not docPr.get("descr"):
                    self.issues.append(Issue("IMAGE-ALT", "ERROR", para.text,
                                             "Missing alt text"))
                    continue

                descr = (docPr.get("descr") or "").strip()
                name = (docPr.get("name") or "").strip()
                sample = f"{name} {descr}".lower()
                if any(token in sample for token in ("table", "tabular", "spreadsheet", "grid")):
                    self.issues.append(Issue(
                        "TABLE-AS-IMAGE",
                        "ERROR",
                        para.text[:500],
                        "Table appears to be inserted as an image instead of a real table.",
                        fix_applied=False,
                    ))

    # ================= EQUATIONS =================
    def _check_equations(self):
        for para in self.doc.paragraphs:
            if "=" in para.text:
                for run in para.runs:
                    if run.bold or run.italic:
                        self.issues.append(Issue("EQ-FORMAT", "ERROR",
                                                 para.text, "Bad equation format"))

    # ================= META =================
    def _check_metadata(self):
        if not self.doc.core_properties.title:
            self.issues.append(Issue("META", "ERROR", "Doc", "Missing title"))

    # ================= OTHER =================
    def _check_color_contrast(self):
        pass

    def _check_page_breaks(self):
        # Report likely unnecessary explicit page/section breaks.
        # To preserve originality, this check does not auto-remove breaks.
        for para in self._iter_all_paragraphs():
            p_el = para._p
            br_nodes = p_el.findall(".//" + qn("w:br"))
            page_breaks = [b for b in br_nodes if (b.get(qn("w:type")) or "") in ("page", "column")]
            sect_nodes = p_el.findall(".//" + qn("w:sectPr"))

            if len(page_breaks) > 1:
                self.issues.append(Issue(
                    "PAGEBREAK-MULTI",
                    "WARNING",
                    para.text[:500],
                    "Multiple explicit page/column breaks in one paragraph; verify necessity.",
                    fix_applied=False,
                ))

            if page_breaks and not (para.text or "").strip():
                self.issues.append(Issue(
                    "PAGEBREAK-EMPTY-PARA",
                    "WARNING",
                    "Paragraph break marker",
                    "Explicit page/column break on an otherwise empty paragraph; verify necessity.",
                    fix_applied=False,
                ))

            if len(sect_nodes) > 1:
                self.issues.append(Issue(
                    "SECTIONBREAK-MULTI",
                    "WARNING",
                    para.text[:500],
                    "Multiple section-break definitions in one paragraph; verify necessity.",
                    fix_applied=False,
                ))

    # ================= SUMMARY =================
    def _print_summary(self):
        errors = [i for i in self.issues if i.severity == "ERROR"]
        warnings = [i for i in self.issues if i.severity == "WARNING"]

        print("\n--- SUMMARY ---")
        print(f"Total: {len(self.issues)}")
        print(f"Errors: {len(errors)}")
        print(f"Warnings: {len(warnings)}\n")

        for issue in self.issues:
            self._safe_print(issue)

    @staticmethod
    def _safe_print(value):
        text = str(value)
        encoding = getattr(sys.stdout, "encoding", None) or "utf-8"
        try:
            text.encode(encoding)
            print(text)
        except UnicodeEncodeError:
            print(text.encode(encoding, errors="replace").decode(encoding, errors="replace"))
